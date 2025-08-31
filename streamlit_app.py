import json
from typing import Optional, List, Dict, Any, Union
from dataclasses import dataclass

import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

import torch
from transformers import pipeline

# -----------------------------
#  Small clinical “tools”
# -----------------------------

NORMALS = {
    "hemoglobin": (12.0, 16.0),
    "platelets": (150, 450),
    "wbc": (4.0, 11.0),
    "ca 15-3": (0, 30),
}

def lab_flagger(labs_json: str) -> Dict[str, Any]:
    try:
        labs = json.loads(labs_json)
        if not isinstance(labs, list):
            return {"flags": [{"error": "Input must be a JSON list of lab dicts."}]}
        flags = []
        for item in labs:
            test = str(item.get("test", "")).lower()
            val = item.get("value")
            rng = NORMALS.get(test)
            if rng and val is not None:
                lo, hi = rng
                if val < lo or val > hi:
                    flags.append({"test": test, "value": val, "normal_range": [lo, hi]})
        return {"flags": flags}
    except Exception as e:
        return {"flags": [{"error": str(e)}]}

def eligibility_check(criteria_json: str, patient_json: str) -> Dict[str, Any]:
    try:
        crit = json.loads(criteria_json)
        pat  = json.loads(patient_json)

        age = (
            pat.get("demographics", {}).get("age_years")
            or pat.get("demographics", {}).get("age")
        )
        ok = True
        reasons = []

        if crit.get("age_min") is not None and (age is None or age < crit["age_min"]):
            ok, reasons = False, reasons + [f"Age < {crit['age_min']}"]
        if crit.get("age_max") is not None and (age is None or age > crit["age_max"]):
            ok, reasons = False, reasons + [f"Age > {crit['age_max']}"]

        req = set([c.lower() for c in crit.get("required_active_conditions", [])])
        if req:
            actives = {
                (c.get("description") or "").lower()
                for c in pat.get("conditions", [])
                if c.get("active") is True
            }
            missing = req - actives
            if missing:
                ok, reasons = False, reasons + [f"Missing active conditions: {sorted(list(missing))}"]

        excl = set([c.lower() for c in crit.get("excluded_conditions", [])])
        if excl:
            present = {
                (c.get("description") or "").lower()
                for c in pat.get("conditions", [])
            } & excl
            if present:
                ok, reasons = False, reasons + [f"Has excluded conditions: {sorted(list(present))}"]

        return {"eligible": ok, "reasons": reasons}
    except Exception as e:
        return {"eligible": None, "reasons": [str(e)]}

# -----------------------------
#  LLM Setup (FLAN-T5)
# -----------------------------

@st.cache_resource(show_spinner=True)
def load_pipe(model_name: str = "google/flan-t5-base"):
    device = 0 if torch.cuda.is_available() else -1  # Streamlit Cloud: CPU -> -1
    return pipeline("text2text-generation", model=model_name, device=device)

def build_patient_prompt(data: Dict[str, Any]) -> str:
    # clip list lengths to keep prompt small
    conds = (data.get("conditions") or [])[:5]
    d = {
        "patient_id": data.get("patient_id"),
        "demographics": data.get("demographics", {}),
        "conditions": conds,
    }
    exemplar = (
        "Example style:\n"
        "Hi there, this is a quick update based on your record. You have active high blood pressure and prediabetes. "
        "You’re a 45-year-old female. Keep taking medicines as directed and try to stay active most days. "
        "Please schedule your routine checkup in the next few months. "
        "Contact your care team if symptoms change."
    )
    return f"""
You are a medical assistant writing a brief, friendly note to a patient.
Use ONLY the facts supplied. Do not invent diagnoses or dates.
Plain language (grade 6–8). Write 4–6 sentences.

Facts (JSON):
{json.dumps(d, indent=2)}

{exemplar}

Now write the patient message in the same style (no headings, no lists):
""".strip()

def build_trial_prompt(data: Dict[str, Any]) -> str:
    d = dict(data)
    # clip sequences
    for k in ["visits", "medications", "labs", "adverse_events"]:
        if isinstance(d.get(k), list):
            d[k] = d[k][:5]
    return f"""
You are a clinical documentation assistant. Summarize these oncology/trial facts into 4–6 sentences.
Be precise, avoid speculation; include stage, current therapy, recent visits/labs/AEs if present, and next planned visit.

Facts (JSON):
{json.dumps(d, indent=2)}

Write the concise clinical summary (no headings/bullets):
""".strip()

def detect_kind(payload: Dict[str, Any]) -> str:
    keys = set(payload.keys())
    if {"baseline", "visits", "medications", "labs", "adverse_events"} & keys:
        return "oncology"
    return "general"

def clean_narrative(txt: str) -> str:
    txt = txt.strip()
    # Guarantee the required closing line for patient-facing messages
    if not txt.endswith("Contact your care team if symptoms change."):
        if not txt.endswith("."):
            txt += "."
        txt += " Contact your care team if symptoms change."
    return txt

def generate_narrative(pipe, json_str: str) -> str:
    try:
        payload = json.loads(json_str)

        # If user pasted a list by mistake, route to lab flagger
        if isinstance(payload, list):
            flags = lab_flagger(json_str)
            return json.dumps(flags, indent=2)

        kind = detect_kind(payload)
        prompt = build_trial_prompt(payload) if kind == "oncology" else build_patient_prompt(payload)

        out = pipe(
            prompt,
            max_new_tokens=200,
            no_repeat_ngram_size=3,
            length_penalty=0.9,
            num_beams=4,
            do_sample=False,
        )
        text = out[0]["generated_text"].strip()
        if kind == "general":
            text = clean_narrative(text)
        return text
    except Exception as e:
        return f"[Narrative error] {e}"

# -----------------------------
#  DOCX helpers
# -----------------------------

def narrative_to_docx(narrative: str, title: str = "Patient Narrative") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in narrative.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# -----------------------------
#  Streamlit UI
# -----------------------------

st.set_page_config(page_title="Clinical Trial Ops (FLAN-T5)", page_icon="🩺", layout="centered")
st.title("🩺 Clinical Trial Ops Assistant (FLAN-T5)")

with st.sidebar:
    st.subheader("Model")
    model_choice = st.selectbox(
        "Choose model (Cloud: start with base)",
        ["google/flan-t5-base", "google/flan-t5-small", "google/flan-t5-large"],
        index=0,
        help="Streamlit Cloud has limited RAM/CPU; 'base' is a safe default."
    )
    pipe = load_pipe(model_choice)
    st.caption("Loaded. (Cloud runs CPU; local machines with CUDA use GPU automatically.)")

tabs = st.tabs(["Narrative", "Eligibility", "Lab flags", "Batch"])

# ---- Narrative tab ----
with tabs[0]:
    st.markdown("Paste a **patient JSON** (demographics + active conditions).")
    example_payload = {
        "patient_id": "X001",
        "demographics": {"age_years": 55, "gender": "M"},
        "conditions": [{"description": "Diabetes", "active": True}],
    }
    payload_in = st.text_area("Patient JSON", value=json.dumps(example_payload, indent=2), height=220)
    if st.button("Generate narrative"):
        output = generate_narrative(pipe, payload_in)
        st.write(output)
        if not output.startswith("[Narrative error]") and output.strip():
            docx_bytes = narrative_to_docx(output, title=f"Narrative {example_payload['patient_id']}")
            st.download_button("Download DOCX", data=docx_bytes, file_name="narrative.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---- Eligibility tab ----
with tabs[1]:
    st.markdown("Provide **criteria JSON** and **patient JSON**.")
    crit_ex = {"age_min": 18, "age_max": 70, "required_active_conditions": ["Hypertension"], "excluded_conditions": ["Pregnancy"]}
    pat_ex  = {"patient_id": "1234", "demographics": {"age_years": 45}, "conditions": [{"description": "Hypertension", "active": True}]}
    cjson = st.text_area("Criteria JSON", value=json.dumps(crit_ex, indent=2), height=160)
    pjson = st.text_area("Patient JSON",  value=json.dumps(pat_ex,  indent=2), height=160)
    if st.button("Check eligibility"):
        res = eligibility_check(cjson, pjson)
        st.json(res)

# ---- Lab flags tab ----
with tabs[2]:
    st.markdown("Paste a **labs JSON list** (e.g., `[{'test':'WBC','value':12.3,'unit':'10^9/L'}]`).")
    labs_ex = [
        {"test":"Hemoglobin","value":10.8,"unit":"g/dL"},
        {"test":"Platelets","value":220,"unit":"10^9/L"},
        {"test":"WBC","value":12.3,"unit":"10^9/L"},
    ]
    labs_in = st.text_area("Labs JSON list", value=json.dumps(labs_ex, indent=2), height=220)
    if st.button("Flag labs"):
        flags = lab_flagger(labs_in)
        st.json(flags)
        # Simple chart for numeric values
        try:
            data = json.loads(labs_in)
            if isinstance(data, list) and data:
                dfv = pd.DataFrame([{"test": d.get("test"), "value": d.get("value")} for d in data if "value" in d])
                if not dfv.empty:
                    st.bar_chart(dfv.set_index("test"))
        except Exception:
            pass

# ---- Batch tab ----
with tabs[3]:
    st.markdown("""
Upload a CSV with columns you have (any subset is okay):
- **patient_id**
- **payload_json**  (required for narrative)
- **criteria_json** (optional)
- **patient_json**  (optional; if missing, we’ll reuse payload_json)
- **labs_json**     (optional)
""")
    up = st.file_uploader("Upload CSV", type=["csv"])
    if up is not None:
        df = pd.read_csv(up)
        out_rows = []
        for _, row in df.iterrows():
            pid = row.get("patient_id", "")
            # Narrative
            narrative = ""
            pj = row.get("payload_json", "")
            if isinstance(pj, str) and pj.strip():
                narrative = generate_narrative(pipe, pj)
            # Eligibility
            eligible, reasons = None, []
            cj = row.get("criteria_json", "")
            pj2 = row.get("patient_json", "") or pj
            if isinstance(cj, str) and cj.strip() and isinstance(pj2, str) and pj2.strip():
                er = eligibility_check(cj, pj2)
                eligible, reasons = er.get("eligible"), er.get("reasons", [])
            # Labs
            lf = ""
            lj = row.get("labs_json", "")
            if isinstance(lj, str) and lj.strip():
                lf = json.dumps(lab_flagger(lj), indent=2)
            out_rows.append({
                "patient_id": pid,
                "narrative": narrative,
                "eligible": eligible,
                "reasons": reasons,
                "lab_flags": lf
            })
        out_df = pd.DataFrame(out_rows)
        st.dataframe(out_df, use_container_width=True)
        csv_bytes = out_df.to_csv(index=False).encode("utf-8")
        st.download_button("Download results CSV", data=csv_bytes, file_name="batch_results.csv", mime="text/csv")
